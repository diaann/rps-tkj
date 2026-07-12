#!/usr/bin/env python3
"""
docx_extractor.py

baca dokumen rps (.docx) pakai python-docx, cuma baca isi tabel (bukan paragraf biasa,
jadi cover otomatis skip). cari tabel identitas lewat header "mata kuliah"+"kode"+"semester".
output json-nya dipakai rpsDocxParser.js buat dirakit jadi rpsData.
"""

import sys
import json
import re
import warnings

warnings.filterwarnings("ignore")

for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

try:
    import docx
    from docx.oxml.ns import qn
except ImportError:
    print(json.dumps({"error": "python_docx_not_installed"}))
    sys.exit(1)

from rps_common import (
    clean,
    clean_multiline,
    extract_cp_tree,
    find_authoritative_subcpmk_owners,
    extract_assessment_rows,
    extract_narrative_sections,
    parse_embedded_assessment,
)


# ---------------------------------------------------------------------------
# baca tabel python-docx
# ---------------------------------------------------------------------------

def dedupe_row(row):
    # merged cell bikin python-docx return Cell object yg sama berkali-kali per kolom.
    # cek pakai cell._tc (identity), bukan teks, biar 2 sel kosong yg beda kolom ga ketuker.
    result = []
    last_tc = None
    for cell in row.cells:
        if cell._tc is last_tc:
            continue
        result.append(clean_multiline(cell.text))
        last_tc = cell._tc
    return result

# menerapkan dedupe_row ke setiap baris sebuah tabel word
# menghasilkan representasi tabel sbg list of list string
def table_to_rows(table):
    return [dedupe_row(row) for row in table.rows]


def expand_multiline_cells(rows):
    # 1 sel bisa isinya banyak butir (\n = batas paragraf beneran di word).
    # pecah jadi baris sendiri per butir sblm ke extract_narrative_sections.
    expanded = []
    for row in rows:
        if len(row) < 2:
            expanded.append(row)
            continue
        head = row[:-1]
        lines = [l for l in row[-1].split("\n") if l.strip()]
        if len(lines) <= 1:
            expanded.append(row)
            continue
        expanded.append(list(head) + [lines[0]])
        blank_head = [""] * len(head)
        for line in lines[1:]:
            expanded.append(blank_head + [line])
    return expanded

#  mencari tabel mana yg merupaka tabel identitas rps
def find_identity_table_index(tables_as_rows):
    for idx, rows in enumerate(tables_as_rows):
        for row in rows[:6]:
            joined = " ".join(row)
            if (
                re.search(r"MATA\s*KULIAH", joined, re.I)
                and re.search(r"\bKODE\b", joined, re.I)
                and re.search(r"SEMESTER", joined, re.I)
            ):
                return idx
    return None


# ---------------------------------------------------------------------------
# identitas & otorisasi. baca sel per posisi, struktur tabelnya eksplisit
# ---------------------------------------------------------------------------

ROMAN_NUMERAL_VALUES = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}


def roman_to_int(text):
    # angka romawi -> int (buat kolom semester yg kadang ditulis romawi). invalid -> none.
    s = text.strip().upper()
    if not s or not re.fullmatch(r"[IVXLCDM]+", s):
        return None
    total = 0
    prev_value = 0
    for ch in reversed(s):
        value = ROMAN_NUMERAL_VALUES[ch]
        if value < prev_value:
            total -= value
        else:
            total += value
            prev_value = value
    return total if 1 <= total <= 14 else None


def parse_semester_value(text):
    # nomor semester dari 1 sel, terima angka biasa ("4") atau romawi ("IV")
    c_clean = text.strip()
    if re.fullmatch(r"\d{1,2}", c_clean):
        return c_clean
    roman_value = roman_to_int(c_clean)
    if roman_value is not None:
        return str(roman_value)
    return None

# fungsi untuk mencari baris header yg ada kata mata kuliah, kode, dan semester
def extract_identity_docx(rows):
    header_idx = None
    for i, row in enumerate(rows):
        joined = " ".join(row)
        if (
            re.search(r"MATA\s*KULIAH", joined, re.I)
            and re.search(r"\bKODE\b", joined, re.I)
            and re.search(r"SEMESTER", joined, re.I)
        ):
            header_idx = i
            break
    if header_idx is None:
        return {}

    value_row = None
    for j in range(header_idx + 1, min(header_idx + 5, len(rows))):
        row = rows[j]
        if re.search(r"OTORISASI", " ".join(row), re.I):
            break
        if any(c.strip() for c in row):
            value_row = row
            break
    if not value_row:
        return {}

    nama_mk = value_row[0] if len(value_row) > 0 else ""
    kode_mk = value_row[1] if len(value_row) > 1 else ""
    rumpun_mk = value_row[2] if len(value_row) > 2 else ""
    rest = value_row[3:]
    joined_rest = " ".join(rest)

    bobot_match = re.search(r"T\s*[=:]?\s*(\d*)\s*P\s*[=:]?\s*(\d*)", joined_rest, re.I)
    sks_teori = bobot_match.group(1) if bobot_match else ""
    sks_praktikum = bobot_match.group(2) if bobot_match else ""

    semester = ""
    tanggal = ""
    for c in rest:
        c_clean = c.strip()
        if not c_clean:
            continue
        if re.fullmatch(r"T\s*[=:]?\s*\d*", c_clean, re.I) or re.fullmatch(r"P\s*[=:]?\s*\d*", c_clean, re.I):
            continue
        if not semester:
            semester_value = parse_semester_value(c_clean)
            if semester_value is not None:
                semester = semester_value
                continue
        if re.search(r"\d{4}", c_clean) or re.search(
            r"(januari|februari|maret|april|mei|juni|juli|agustus|september|oktober|november|desember|"
            r"january|february|march|may|june|july|august|october|december)",
            c_clean, re.I,
        ):
            tanggal = c_clean

    if re.fullmatch(r"[\?\-\s]*", kode_mk or ""):
        kode_mk = ""

    return {
        "nama_mk": clean(nama_mk),
        "kode_mk": clean(kode_mk),
        "rumpun_mk": clean(rumpun_mk),
        "sks_teori": sks_teori,
        "sks_praktikum": sks_praktikum,
        "semester": semester,
        "tanggal_penyusunan_raw": clean(tanggal),
    }


# cari baris Otorisasi (Pengembang RPS/Koordinator RMK/Ketua Prodi) & ambil nilainya per kolom
def extract_otorisasi_docx(rows):
    header_idx = None
    for i, row in enumerate(rows):
        joined = " ".join(row)
        if re.search(r"OTORISASI", joined, re.I) and re.search(r"Pengembang", joined, re.I):
            header_idx = i
            break
    if header_idx is None:
        return {}

    header_row = rows[header_idx]
    # cari posisi tiap label di header, biar urutan kolom bener meski ada label ekstra
    pengembang_idx = koordinator_idx = ketua_idx = None
    for i, c in enumerate(header_row):
        if re.fullmatch(r"Pengembang(?:\s+RPS)?", c.strip(), re.I):
            pengembang_idx = i
        elif re.fullmatch(r"Koordinator(?:\s+RMK)?", c.strip(), re.I):
            koordinator_idx = i
        elif re.fullmatch(r"Ketua(?:\s+PRODI)?", c.strip(), re.I):
            ketua_idx = i

    value_row = None
    for j in range(header_idx + 1, min(header_idx + 4, len(rows))):
        row = rows[j]
        if re.search(r"Capaian\s+Pembelajaran|CPL", " ".join(row), re.I):
            break
        if any(c.strip() for c in row):
            value_row = row
            break
    if not value_row:
        return {}

    def get(idx):
        return clean(value_row[idx]) if idx is not None and idx < len(value_row) else ""

    return {
        "pengembang_rps": get(pengembang_idx),
        "koordinator_rmk": get(koordinator_idx),
        "ketua_prodi": get(ketua_idx),
    }


# ---------------------------------------------------------------------------
# rencana pembelajaran mingguan. cari header "pekan"+"indikator"+"teknik",
# lalu ambil semua baris data sesudahnya
# ---------------------------------------------------------------------------


# baris minggu evaluasi (uts/uas) = 1 sel gabungan panjang, bukan baris sub-cpmk asli. skip.
EVALUATION_WEEK_PATTERN = re.compile(
    r"ujian\s+tengah\s+semester|ujian\s+akhir\s+semester|\bUTS\b|\bUAS\b|"
    r"evaluasi\s+tengah\s+semester|evaluasi\s+akhir\s+semester",
    re.I,
)


# baca tabel Rencana Pembelajaran Mingguan (1 baris = 1 pekan), skip minggu UTS-UAS
def extract_weekly_plan_docx(all_rows):
    weekly = []
    seen_header = False
    for row in all_rows:
        joined = " ".join(c for c in row if c.strip())
        if not joined:
            continue

        if not seen_header:
            if re.search(r"Pekan|\bMg\s*Ke", joined, re.I) and re.search(r"Indikator|Sub-?CPMK", joined, re.I):
                seen_header = True
            continue

        non_empty_cells = [c.strip() for c in row if c.strip()]
        if non_empty_cells and all(re.fullmatch(r"\(\d{1,2}\)", c) for c in non_empty_cells):
            continue

        first_cell = row[0].strip() if row else ""
        pekan_match = re.fullmatch(r"\(?(\d{1,2})\)?(?:\s*[-\u2013]\s*\(?(\d{1,2})\)?)?", first_cell)
        if not pekan_match:
            continue

        # index kolom tetap: 0=pekan, 1=sub-cpmk, 2=indikator, 3=teknik&kriteria, 4=luring, 5=daring, 6=materi, 7=bobot
        def get(i):
            return row[i] if i < len(row) else ""

        if EVALUATION_WEEK_PATTERN.search(get(1)):
            continue

        pekan_awal = pekan_match.group(1)
        pekan_akhir = pekan_match.group(2) or pekan_awal
        bobot_raw = get(7).strip()
        bobot_raw_clean = re.sub(r"\s+", "", bobot_raw)
        bobot_match = re.fullmatch(r"(\d{1,3}(?:[.,]\d+)?)\s*%?", bobot_raw_clean)
        bobot = bobot_match.group(1) if bobot_match else ""

        weekly.append({
            "pekan_awal": pekan_awal,
            "pekan_akhir": pekan_akhir,
            "kemampuan_akhir": get(1),
            "indikator": get(2),
            "teknik_kriteria": get(3),
            "metode_luring": get(4),
            "metode_daring": get(5),
            "materi": get(6),
            "bobot": bobot,
        })
    return weekly


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

STOP_SECTION_PATTERNS = [
    r"TABEL\s+RUBRIK",
    r"JUDUL\s+TUGAS",
    r"RANCANGAN\s+TUGAS",
    r"Aspek\s+Penilaian",
    r"Rubrik\s+Penilaian",
]


# cek baris ini "tabel lampiran" (rubrik/tugas) yg bukan bagian RPS inti -> berhenti baca di sini
def is_stop_section(row):
    joined = " ".join(c for c in row if c.strip())
    return any(re.search(p, joined, re.I) for p in STOP_SECTION_PATTERNS)


# orkestrator utama: buka .docx, cari tabel identitas sbg jangkar, lalu panggil semua fungsi ekstraksi di atas
def extract(docx_path):
    document = docx.Document(docx_path)
    tables_as_rows = [table_to_rows(t) for t in document.tables]

    anchor_idx = find_identity_table_index(tables_as_rows)
    if anchor_idx is None:
        return {"error": "identity_table_not_found"}

    identity_rows = tables_as_rows[anchor_idx]
    identity = extract_identity_docx(identity_rows)
    otorisasi = extract_otorisasi_docx(identity_rows)

    # identity/otorisasi selalu di tabel pertama (anchor), tapi cp-tree/asesmen/naratif/mingguan
    # kadang lanjut ke tabel word berikutnya. gabung semua sampai ketemu stop section
    combined_table_indices = [anchor_idx]
    combined_rows = list(identity_rows)
    for idx in range(anchor_idx + 1, len(tables_as_rows)):
        rows = tables_as_rows[idx]
        if rows and is_stop_section(rows[0]):
            break
        combined_table_indices.append(idx)
        combined_rows.extend(rows)

    # cari nested table di dalam sel (tabel "korelasi cp & asesmen" biasanya di sini).
    # simpen objek cell._tc-nya sendiri di set (bukan id()) biar identity check-nya valid.
    nested_tables_rows = []
    seen_tcs = set()
    for idx in combined_table_indices:
        for row in document.tables[idx].rows:
            for cell in row.cells:
                if cell._tc in seen_tcs:
                    continue
                seen_tcs.add(cell._tc)
                for nested in cell.tables:
                    nested_tables_rows.append(table_to_rows(nested))

    # ambil range baris cp-tree: dari label "capaian pembelajaran" pertama sampai "korelasi cp & asesmen"
    cp_start_idx = None
    cp_end_idx = None
    for i, row in enumerate(combined_rows):
        joined = " ".join(row)
        if cp_start_idx is None and (
            re.match(r"^Capaian\s+Pembelajaran", row[0], re.I)
            or re.search(r"CPL[\s-]?PRODI", joined, re.I)
        ):
            cp_start_idx = i
        if cp_start_idx is not None and re.search(r"Korelasi\s+(?:antara\s+)?CP\s+dan\s+(?:[Aa]sesmen|[Aa]ssessmen|[Aa]ssessment)", joined, re.I):
            cp_end_idx = i
            break
    if cp_start_idx is None:
        cp_rows_raw = []
    else:
        cp_rows_raw = combined_rows[cp_start_idx:(cp_end_idx + 1) if cp_end_idx is not None else len(combined_rows)]

    cp_rows_stripped = []
    for row in cp_rows_raw:
        if len(row) > 1 and re.match(r"^Capaian\s+Pembelajaran", row[0], re.I):
            cp_rows_stripped.append(row[1:])
        else:
            cp_rows_stripped.append(row)

    cp_all_tables = [cp_rows_stripped] + nested_tables_rows
    # pemetaan cpmk<->subcpmk yg jelas (authoritative_owners) dicari di combined_rows yg lebih luas,
    # soalnya suka ada di tabel rencana mingguan yg udah kepotong dari cp_rows_raw
    authoritative_owners = find_authoritative_subcpmk_owners([combined_rows] + nested_tables_rows)
    cpl_list, cpmk_list, subcpmk_list = extract_cp_tree(cp_all_tables, authoritative_owners=authoritative_owners)

    # weekly dihitung duluan di sini krn dipakai juga sbg fallback formatif/sumatif di bawah
    weekly = extract_weekly_plan_docx(combined_rows)

    # dipanggil 2x terpisah (bukan digabung 1 list): trigger text "korelasi cp & asesmen" cuma
    # ada di combined_rows, sedangkan isi kuis/tugas/ujian beneran sering di nested table.
    # nested table di sini udah pasti scope-nya korelasi, jadi assume_in_section=True langsung.
    assessment_rows = {
        **extract_assessment_rows(nested_tables_rows, assume_in_section=True),
        **extract_assessment_rows([combined_rows]),
    }
    for idx, sub in enumerate(subcpmk_list):
        info = assessment_rows.get(sub.get("raw_number", sub["global_number"])) or {}
        has_korelasi_data = bool(
            info.get("formatif") or info.get("kuis") or info.get("tugas")
            or info.get("ujian") or info.get("pjbl") or info.get("presentasi") or info.get("lainnya")
        )

        # fallback dihitung selalu, bukan cuma pas has_korelasi_data == false, soalnya tabel korelasi
        # kadang cuma punya sebagian bucket (mis. presentasi cuma disebut di teks bebas mingguan)
        week_detail = weekly[idx] if idx < len(weekly) else None
        parsed = parse_embedded_assessment((week_detail or {}).get("teknik_kriteria", ""))

        if has_korelasi_data:
            sub["bobot"] = info.get("bobot_total", "")
            sub["formatif_nama"] = info.get("formatif") or parsed["formatif"]
            sub["formatif_bobot"] = info.get("formatif_bobot") or parsed["formatif_bobot"]
            sub["kuis"] = info.get("kuis") or parsed["kuis"]
            sub["tugas"] = info.get("tugas") or parsed["tugas"]
            sub["ujian"] = info.get("ujian") or parsed["ujian"]
            sub["pjbl"] = info.get("pjbl") or parsed["pjbl"]
            sub["presentasi"] = info.get("presentasi") or parsed["presentasi"]
            sub["lainnya"] = info.get("lainnya") or parsed["lainnya"] or []
            sub["bentuk_pembelajaran"] = info.get("bentuk_pembelajaran", "")
            sub["metode_pembelajaran"] = info.get("metode_pembelajaran", "")
        else:
            # ga ada tabel korelasi sama sekali -> rincian formatif/sumatif ditulis bebas
            # di kolom teknik&kriteria mingguan. sub-cpmk ke-n dipasangin ke baris mingguan ke-n.
            if parsed["formatif"] or parsed["formatif_bobot"] or parsed["kuis"] or parsed["tugas"] or parsed["ujian"] or parsed["pjbl"] or parsed["presentasi"] or parsed["lainnya"]:
                sub["formatif_nama"] = parsed["formatif"]
                sub["formatif_bobot"] = parsed["formatif_bobot"]
                sub["kuis"] = parsed["kuis"]
                sub["tugas"] = parsed["tugas"]
                sub["ujian"] = parsed["ujian"]
                sub["pjbl"] = parsed["pjbl"]
                sub["presentasi"] = parsed["presentasi"]
                sub["lainnya"] = parsed["lainnya"]

    # sama kayak cp-tree: ambil range kontinu dari label naratif pertama sampai header
    # rencana mingguan, bukan filter per baris (biar baris lanjutan yg ga ada labelnya ga kebuang)
    narrative_start_idx = None
    narrative_end_idx = None
    for i, row in enumerate(combined_rows):
        if narrative_start_idx is None and re.match(
            r"^(Deskripsi\s+Singkat|Materi\s+Kajian|Bahan\s+Kajian|Pustaka|Utama|Pendukung|"
            r"Dosen(?:\s+Pengampu)?|Mata\s*kuliah\s+syarat|Matakuliah\s+syarat)",
            row[0], re.I,
        ):
            narrative_start_idx = i
        if narrative_start_idx is not None and (
            re.match(r"^Pekan\b|^Mg\s*Ke", row[0], re.I) or row[0].strip() == "(1)"
        ):
            narrative_end_idx = i
            break
    if narrative_start_idx is None:
        narrative_rows_raw = []
    else:
        narrative_rows_raw = combined_rows[
            narrative_start_idx:(narrative_end_idx if narrative_end_idx is not None else len(combined_rows))
        ]
    narrative_rows = expand_multiline_cells(narrative_rows_raw)
    text_sections = extract_narrative_sections([narrative_rows])

    full_text = "\n".join(p.text for p in document.paragraphs)

    return {
        "identity": identity,
        "otorisasi": otorisasi,
        "cpl": cpl_list,
        "cpmk": cpmk_list,
        "subcpmk": subcpmk_list,
        "weekly": weekly,
        **text_sections,
        "raw_text": full_text,
    }


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "no_docx_path_given"}))
        sys.exit(1)
    try:
        result = extract(sys.argv[1])
        print(json.dumps(result, ensure_ascii=True))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)
